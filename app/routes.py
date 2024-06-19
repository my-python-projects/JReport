from app import app
from flask import render_template, request, send_file, jsonify
from docx import Document
from datetime import datetime
import os


@app.route('/')
def index():
    app.logger.info('Acesso à página inicial')

    # Teste se já está logado
    isLogin = True
    # teste com condição if ternario
    return render_template('index.html' if isLogin else 'login.html')

@app.route('/login')
def login():
    app.logger.info('Acesso à página Login')
    return render_template('login.html')


@app.route('/register')
def register():
    app.logger.info('Acesso à página de Cadastro')
    return render_template('register.html')

@app.route('/processar_formulario', methods=['POST'])
def processar_formulario():

    try:

        nome = request.form['nome']
        equipamento = request.form['equipamento']
        inicio = request.form['inicio']
        final = request.form['final']
        estoque_utilizado = request.form['estoqueUtilizado']
        removido_voltou = request.form['removidoVoltou']
        motivo_remocao = request.form.get('motivoRemocao', '')

        # Create a Word document and write the data
        document = Document('./template.docx')

        document.add_paragraph(f'Nome: {nome}')
        document.add_paragraph(f'Equipamento: {equipamento}')
        document.add_paragraph(f'Período: {inicio} - {final}')
        document.add_paragraph(f'Estoque Utilizado: {estoque_utilizado}')

        document.add_paragraph('Algo foi removido e voltou ao estoque? ' + ('Sim' if removido_voltou == 'sim' else 'Não'))

        if removido_voltou == 'sim':
            document.add_paragraph(f'Motivo da Remoção: {motivo_remocao}')

        # Get current date
        data_atual = datetime.now()

        # Format the date as a string
        formato_data = "%d-%m-%Y"
        data_formatada = data_atual.strftime(formato_data)

        caminho = './relatorios/'

        # Check that the directory does not exist before creating
        if not os.path.exists(caminho):
            os.makedirs(caminho)
            print(f'The folder at {caminho} was created successfully.')

        documento = "{}{}.docx".format(caminho, data_formatada)

        document.save(documento)

        #send_file(documento, as_attachment=True)
        return jsonify(success=True)
    
    except Exception as e:
        print(f'Erro: {str(e)}')
        return jsonify(success=False)

